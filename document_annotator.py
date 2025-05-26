import streamlit as st
import docx
import io
from openai import OpenAI
import os
import json
from datetime import datetime
import tempfile
import fitz  # PyMuPDF
from PIL import Image
import numpy as np
import base64
from bs4 import BeautifulSoup
import shutil

# 设置页面配置
st.set_page_config(layout="wide")

# 创建临时目录用于存储转换后的文件
TEMP_DIR = "temp_files"
if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR)

# 初始化session state
if "annotations" not in st.session_state:
    st.session_state.annotations = []
if "current_file_content" not in st.session_state:
    st.session_state.current_file_content = ""
if "file_name" not in st.session_state:
    st.session_state.file_name = ""
if "selected_text" not in st.session_state:
    st.session_state.selected_text = ""
if "current_page" not in st.session_state:
    st.session_state.current_page = 0
if "total_pages" not in st.session_state:
    st.session_state.total_pages = 0
if "pdf_document" not in st.session_state:
    st.session_state.pdf_document = None
if "pdf_scale" not in st.session_state:
    st.session_state.pdf_scale = 1.0
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "current_selection" not in st.session_state:
    st.session_state.current_selection = None

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def process_file(uploaded_file):
    if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(uploaded_file)
    return ""

def add_annotation(selected_text, comment, annotation_type, page_num=None):
    st.session_state.annotations.append({
        "text": selected_text,
        "comment": comment,
        "type": annotation_type,
        "page": page_num,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "file_name": st.session_state.file_name
    })

def export_annotations():
    if not st.session_state.annotations:
        return "No annotations to export."
    
    notes = "# 文档标注笔记\n\n"
    
    # 按类型分组导出
    types = ["重要内容", "问题", "评论", "待办事项"]
    for type_name in types:
        type_annotations = [ann for ann in st.session_state.annotations if ann["type"] == type_name]
        if type_annotations:
            notes += f"## {type_name}\n\n"
            for ann in type_annotations:
                page_info = f"[第{ann['page']}页] " if ann.get('page') is not None else ""
                notes += f"### {ann['file_name']} - {page_info}{ann['timestamp']}\n"
                notes += f"**选中文本：** {ann['text']}\n"
                notes += f"**评论：** {ann['comment']}\n\n"
    
    return notes

def create_text_layer(page, zoom=2):
    """创建文本层HTML"""
    blocks = page.get_text("dict")["blocks"]
    text_layer_html = f"""
    <div class="text-layer" style="position: relative; width: 100%; height: 100%;">
    """
    
    # 获取页面尺寸
    page_rect = page.rect
    scale = zoom  # 缩放因子
    
    # 处理每个文本块
    for block in blocks:
        if "lines" in block:
            for line in block["lines"]:
                for span in line["spans"]:
                    # 计算缩放后的位置
                    x0, y0, x1, y1 = span["bbox"]
                    scaled_x = x0 * scale
                    scaled_y = y0 * scale
                    scaled_width = (x1 - x0) * scale
                    scaled_height = (y1 - y0) * scale
                    
                    # 创建文本span元素
                    text_layer_html += f"""
                    <span class="text-block" style="
                        position: absolute;
                        left: {scaled_x}px;
                        top: {scaled_y}px;
                        width: {scaled_width}px;
                        height: {scaled_height}px;
                        font-size: {span['size'] * scale}px;
                        font-family: {span.get('font', 'Arial')};
                        cursor: text;
                        user-select: text;
                        color: transparent;
                        background: transparent;
                        z-index: 1;
                    ">{span['text']}</span>
                    """
    
    text_layer_html += "</div>"
    return text_layer_html

def display_pdf_page(pdf_document, page_num):
    """显示PDF页面并创建文本层"""
    page = pdf_document[page_num]
    zoom = 2  # 2x缩放以获得更好的质量
    
    # 获取页面的图像
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    
    # 创建文本层
    text_layer = create_text_layer(page, zoom)
    
    # 创建包含图像和文本层的HTML
    html_content = f"""
    <div class="pdf-container" style="position: relative; width: 100%;">
        <div class="image-layer" style="position: relative;">
            <img src="data:image/png;base64,{image_to_base64(img)}" style="width: 100%; height: auto;"/>
            {text_layer}
        </div>
    </div>
    <style>
        .pdf-container {{
            position: relative;
            width: 100%;
        }}
        .text-layer {{
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            pointer-events: auto;
        }}
        .text-block {{
            position: absolute;
            pointer-events: auto;
        }}
        .text-block:hover {{
            background: rgba(255, 255, 0, 0.2) !important;
        }}
        ::selection {{
            background: yellow;
            color: black;
        }}
    </style>
    <script>
        document.addEventListener('mouseup', function() {{
            const selectedText = window.getSelection().toString().trim();
            if (selectedText) {{
                window.parent.postMessage({{
                    type: 'text_selected',
                    text: selectedText
                }}, '*');
            }}
        }});
    </script>
    """
    
    # 显示PDF页面和文本层
    st.components.v1.html(html_content, height=pix.height, scrolling=True)
    
    # 返回页面文本（用于搜索功能）
    return page.get_text()

def image_to_base64(image):
    """将PIL图像转换为base64字符串"""
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode()

def handle_text_selection():
    """处理文本选择事件"""
    if st.session_state.get("current_selection"):
        selected_text = st.session_state.current_selection
        if selected_text and isinstance(selected_text, str):
            # 添加选中的文本到聊天历史
            st.session_state.chat_history.append({
                "role": "reference",
                "content": selected_text
            })
            # 清除选择
            st.session_state.current_selection = None

def display_pdf_viewer(pdf_base64):
    """使用 PDF.js 显示 PDF 并启用文本选择"""
    # 创建两列布局：PDF查看器和对话区域
    col_pdf, col_chat = st.columns([2, 1])
    
    with col_pdf:
        # 添加缩放控制按钮
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if st.button("🔍 放大", use_container_width=True):
                st.session_state.pdf_scale = min(2.0, st.session_state.pdf_scale + 0.1)
                st.rerun()
        with col2:
            st.write(f"当前缩放: {st.session_state.pdf_scale:.1f}x")
        with col3:
            if st.button("🔍 缩小", use_container_width=True):
                st.session_state.pdf_scale = max(0.5, st.session_state.pdf_scale - 0.1)
                st.rerun()
        
        viewer_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js"></script>
            <style>
                #viewerContainer {{
                    width: 100%;
                    height: 800px;
                    overflow: auto;
                    background: #404040;
                    text-align: center;
                }}
                .pdfViewer {{
                    margin: 0 auto;
                }}
                .page {{
                    position: relative;
                    margin: 10px auto;
                    border: 1px solid #000;
                    background-color: white;
                }}
                .textLayer {{
                    position: absolute;
                    left: 0;
                    top: 0;
                    right: 0;
                    bottom: 0;
                    overflow: hidden;
                    opacity: 0.2;
                    line-height: 1.0;
                }}
                .textLayer > span {{
                    color: transparent;
                    position: absolute;
                    white-space: pre;
                    cursor: text;
                    transform-origin: 0% 0%;
                }}
                .textLayer ::selection {{
                    background: rgba(0, 0, 255, 0.2);
                }}
                #confirmButton {{
                    position: fixed;
                    display: none;
                    background: #4CAF50;
                    color: white;
                    border: none;
                    padding: 4px 8px;
                    border-radius: 50%;
                    cursor: pointer;
                    font-size: 14px;
                    box-shadow: 0 2px 5px rgba(0,0,0,0.2);
                    z-index: 1000;
                    transition: all 0.3s ease;
                }}
                #confirmButton:hover {{
                    transform: scale(1.1);
                    background: #45a049;
                }}
            </style>
        </head>
        <body>
            <div id="viewerContainer">
                <div id="pdfViewer" class="pdfViewer"></div>
            </div>
            <button id="confirmButton">✅</button>
            <script>
                pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
                
                const pdfData = atob("{pdf_base64}");
                const loadingTask = pdfjsLib.getDocument({{data: pdfData}});
                
                const confirmButton = document.getElementById('confirmButton');
                let selectedText = '';
                
                // 监听文本选择
                document.addEventListener('mouseup', function(e) {{
                    const selection = window.getSelection();
                    selectedText = selection.toString().trim();
                    
                    if (selectedText) {{
                        const range = selection.getRangeAt(0);
                        const rect = range.getBoundingClientRect();
                        const scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                        
                        // 显示确认按钮
                        confirmButton.style.display = 'block';
                        confirmButton.style.left = (rect.right + 5) + 'px';
                        confirmButton.style.top = (rect.top + scrollTop + rect.height/2 - 10) + 'px';
                    }}
                }});
                
                // 点击确认按钮时发送选中的文本
                confirmButton.addEventListener('click', function() {{
                    if (selectedText) {{
                        // 发送选中的文本到 Streamlit
                        window.parent.postMessage({{
                            type: 'text_selected',
                            text: selectedText
                        }}, '*');
                        confirmButton.style.display = 'none';
                        
                        // 清除选择
                        window.getSelection().removeAllRanges();
                    }}
                }});
                
                // 点击其他地方时隐藏确认按钮
                document.addEventListener('mousedown', function(e) {{
                    if (e.target !== confirmButton) {{
                        confirmButton.style.display = 'none';
                    }}
                }});
                
                // 监听来自 Streamlit 的消息
                window.addEventListener('message', function(event) {{
                    if (event.data.type === 'text_selected') {{
                        // 更新 Streamlit 的状态
                        window.parent.Streamlit.setComponentValue(event.data.text);
                    }}
                }}, false);
                
                loadingTask.promise.then(function(pdf) {{
                    const container = document.getElementById('pdfViewer');
                    
                    function renderPage(pageNum) {{
                        pdf.getPage(pageNum).then(function(page) {{
                            const viewport = page.getViewport({{scale: {st.session_state.pdf_scale}}});
                            const canvas = document.createElement('canvas');
                            const context = canvas.getContext('2d');
                            canvas.height = viewport.height;
                            canvas.width = viewport.width;
                            
                            const renderContext = {{
                                canvasContext: context,
                                viewport: viewport
                            }};
                            
                            const renderTask = page.render(renderContext);
                            renderTask.promise.then(function() {{
                                const pageDiv = document.createElement('div');
                                pageDiv.className = 'page';
                                pageDiv.style.width = viewport.width + 'px';
                                pageDiv.style.height = viewport.height + 'px';
                                pageDiv.appendChild(canvas);
                                
                                const textLayerDiv = document.createElement('div');
                                textLayerDiv.className = 'textLayer';
                                textLayerDiv.style.width = viewport.width + 'px';
                                textLayerDiv.style.height = viewport.height + 'px';
                                pageDiv.appendChild(textLayerDiv);
                                
                                container.appendChild(pageDiv);
                                
                                page.getTextContent().then(function(textContent) {{
                                    pdfjsLib.renderTextLayer({{
                                        textContent: textContent,
                                        container: textLayerDiv,
                                        viewport: viewport,
                                        textDivs: []
                                    }});
                                }});
                            }});
                        }});
                    }}
                    
                    // 渲染所有页面
                    for (let i = 1; i <= pdf.numPages; i++) {{
                        renderPage(i);
                    }}
                }});
            </script>
        </body>
        </html>
        """
        
        # 使用 components.v1.html 并获取返回值
        selected_text = st.components.v1.html(viewer_html, height=800, scrolling=True)
        
        # 如果选择了文本，更新 session state
        if selected_text and isinstance(selected_text, str):
            st.session_state.current_selection = selected_text
            handle_text_selection()
    
    # 右侧对话区域
    with col_chat:
        # 添加自定义 CSS 样式
        st.markdown("""
            <style>
                div[data-testid="stChatMessage"] {
                    background: white;
                    border-radius: 10px;
                    padding: 10px;
                    margin: 5px 0;
                }
                div[data-testid="stChatMessage"] p {
                    margin: 0;
                }
                .selected-text-box {
                    background: #F0F7FF !important;
                    border: 1px solid #CCE5FF !important;
                    border-radius: 10px;
                    padding: 15px !important;
                    margin: 10px 0 !important;
                    font-size: 0.95em;
                    line-height: 1.5;
                }
                .selected-text-header {
                    color: #0066CC;
                    font-size: 0.9em;
                    margin-bottom: 10px;
                    display: flex;
                    align-items: center;
                    gap: 5px;
                }
                .selected-text-content {
                    color: #333333;
                    background: white;
                    padding: 10px;
                    border-radius: 5px;
                    border: 1px solid #E6F0FF;
                }
            </style>
        """, unsafe_allow_html=True)
        
        # 创建一个容器来显示所有消息
        chat_container = st.container()
        
        with chat_container:
            # 显示聊天历史
            for message in st.session_state.chat_history:
                if message["role"] == "reference":
                    with st.chat_message("assistant", avatar="📄"):
                        st.markdown(f"""
                            <div class="selected-text-box">
                                <div class="selected-text-header">
                                    <span>📝 已选择的文本段落</span>
                                </div>
                                <div class="selected-text-content">
                                    {message["content"]}
                                </div>
                            </div>
                        """, unsafe_allow_html=True)
                elif message["role"] == "user":
                    with st.chat_message("user"):
                        st.write(message["content"])
                elif message["role"] == "assistant":
                    with st.chat_message("assistant"):
                        st.write(message["content"])
        
        # 用户输入区域（始终保持在底部）
        if prompt := st.chat_input("输入您的问题..."):
            # 检查是否有参考文本
            reference_messages = [msg for msg in st.session_state.chat_history if msg["role"] == "reference"]
            if reference_messages:
                # 获取最近的参考文本
                latest_reference = reference_messages[-1]["content"]
                
                # 添加用户问题到聊天历史
                st.session_state.chat_history.append({
                    "role": "user",
                    "content": prompt
                })
                
                # 显示用户消息
                with st.chat_message("user"):
                    st.write(prompt)
                
                # 调用 OpenAI API
                client = OpenAI(
                    api_key=os.environ.get("OPENAI_API_KEY"),
                    base_url="https://api.openai.com/v1"
                )
                
                # 构建完整的对话历史
                messages = [
                    {
                        "role": "system",
                        "content": "你是一个专业的文档分析助手。请基于用户提供的参考文本，回答用户的问题。回答要简洁、准确、专业。"
                    },
                    {
                        "role": "user",
                        "content": f"参考文本：\n{latest_reference}\n\n请基于这段文本回答我的问题。"
                    }
                ]
                
                # 添加最近的对话历史（从最近的参考文本开始）
                last_reference_index = len(st.session_state.chat_history) - 1
                for i in range(len(st.session_state.chat_history) - 1, -1, -1):
                    if st.session_state.chat_history[i]["role"] == "reference":
                        last_reference_index = i
                        break
                
                recent_messages = st.session_state.chat_history[last_reference_index + 1:]
                messages.extend([
                    msg for msg in recent_messages 
                    if msg["role"] in ["user", "assistant"]
                ])
                
                # 添加当前问题
                messages.append({
                    "role": "user",
                    "content": prompt
                })
                
                # 显示助手正在输入的状态
                with st.chat_message("assistant"):
                    with st.spinner("思考中..."):
                        response = client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=messages,
                            temperature=0.7,
                            max_tokens=500
                        )
                        assistant_response = response.choices[0].message.content
                        st.write(assistant_response)
                
                # 添加助手回答到聊天历史
                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": assistant_response
                })
            else:
                st.warning("请先选择文本")

# 主页面
st.title("文档标注系统")

# 文件上传
uploaded_file = st.file_uploader("上传PDF文档", type=["pdf"])

if uploaded_file:
    st.session_state.file_name = uploaded_file.name
    
    # 处理PDF文件
    if uploaded_file.type == "application/pdf":
        # 将PDF文件转换为base64
        pdf_base64 = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')
        # 显示PDF查看器
        display_pdf_viewer(pdf_base64)

# 清理临时文件
def cleanup():
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)

# 注册清理函数
import atexit
atexit.register(cleanup) 